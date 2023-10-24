import numpy as np
import csv
import os
import pickle
import re
from docx import Document

###-----------------------------------------------------------------------------------------
# Measure the number of human labels in each case

# Filter 300 cases

# Find cases with incorrect labels

# In human label file, the following cases have problems: 110 has different number of labels
# [6, 32, 97, 147, 160, 163, 170, 216, 216, 216, 216, 233, 278, 325, 365, 402, 428, 430, 456, 473, 474, 477]
#
# Load pre_saved key(300)
# Given annotation file of different model and selected key, save detection and correction dict(300 elements)

# Change file and saved file (selected key)



###-----------------------------------------------------------------------------------------
###--------------------------------------------------------------------
# machine ground truth
file_human = 'annotation_data/annotation_human.docx'
llama_file = Document(file_human).paragraphs
llama_lines = [para.text for para in llama_file if para.text.strip()]

detection_file = 'machine_detection'
correction_file = 'machine_correction'


# llama_alpaca result

# file = 'annotation_data/annotation_alpaca.docx'
# file_human = 'annotation_data/annotation_human.docx'
# llama_file = Document(file).paragraphs
# llama_lines = [para.text for para in llama_file if para.text.strip()]
#
# detection_file = 'llama_lora_detection'
# correction_file = 'llama_lora_correction'

###--------------------------------------------------------------------

# no lora result
# file = 'annotation_data/alpaca_no_lora.txt'
#
# with open(file, 'r') as f:
#     lines = f.readlines()
#
# llama_lines = [line for line in lines if line.strip()]
#
#
# detection_file = 'alpaca_no_lora_detection'
# correction_file = 'alpaca_no_lora_correction'


###--------------------------------------------------------------------

# gpt coach result inference result
# file = 'annotation_data/gpt_coach.txt'
#
# with open(file, 'r') as f:
#     lines = f.readlines()
#
# llama_lines = [line for line in lines if line.strip()]
#
# detection_file = 'gpt_coach_detection'
# correction_file = 'gpt_coach_correction'


selected_keys = np.load('eval_keys.npy')

def machine_label(human_lines):
    human_label = {}

    n_labels = {}
    n_case = 0
    wrong_case = []
    for i, line in enumerate(human_lines):
        if line.startswith('Case:'):
            n_case += 1
            human_label[n_case] = []
            n_labels[n_case] = 0

        if line.startswith('机器标注'):
            n_labels[n_case] += 1
            detect_correct = re.split('机器标注：|修改为', line)
            ### try to split line into three parts
            if len(detect_correct) == 3:
                human_label[n_case].append([detect_correct[1], detect_correct[2]])

                ### line has less than 3 elements
            elif len(detect_correct) < 3:

                ### human label is empty, move to next line. Do the same check and split
                if human_lines[i + 1].startswith('错误术语'):

                    detect_correct = re.split('错误术语：|修改为', human_lines[i + 1])
                    if len(detect_correct) == 3:
                        human_label[n_case].append([detect_correct[1], detect_correct[2]])
                    else:
                        print('wrong case of next line')
                        print(n_case)
                        print(detect_correct)
                        wrong_case.append(n_case)
                        human_label[n_case].append(['false', 'false'])


                else:
                    print('wrong case')
                    print(n_case)
                    wrong_case.append(n_case)
                    human_label[n_case].append(['false', 'false'])
                    print(line)
                    print(human_lines[i + 1])

            else:
                print('wrong case')
                print(n_case)
                wrong_case.append(n_case)
                human_label[n_case].append(['false', 'false'])
                print(line)
                print(human_lines[i + 1])
                print('long line', line)
    return human_label

print('llama...')
# llama_machine_label, llama_wrong_case, llama_n_labels = machine_label_process(llama_lines)

llama_case = machine_label(llama_lines)

print(len(llama_case))

detection_label = {}
correction_label = {}



symbols = r'[：。，、（）]|\b错误术语\b|\n'
for label in selected_keys:
    detection_label[label] = []
    correction_label[label] = []
    for pair in llama_case[label]:
        det_label = re.sub(symbols, '', pair[0]).replace(' ', '')
        cor_label = re.sub(symbols, '', pair[1]).replace(' ', '')
        detection_label[label].append(det_label)
        correction_label[label].append(cor_label)

print(len(detection_label))
print(detection_label[4])
lens = 0
for v in detection_label.values():
    lens += len(v)
print(lens)

with open(detection_file, 'wb') as f:
    pickle.dump(detection_label, f)

with open(correction_file, 'wb') as f:
    pickle.dump(correction_label, f)









